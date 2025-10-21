"""
DAX
"""

#!/usr/bin/env python
# coding: utf-8

# In[1]:


import os
import argparse
import base64
import json
import re
import logging
from pathlib import Path

import pandas as pd
import matplotlib.pyplot as plt
import pdfplumber
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Inches

# ----------------------------------------
# Configuración
# ----------------------------------------
load_dotenv()
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)

# Rutas por defecto (ajustar según entorno)


DEFAULT_IMAGE = Path(r"C:\Users\HP\Desktop\PROPUESTAS DE DIPLOMADO\DIPLOMADO EXCEL Y POWER BI APLICADO A LA GESTION DE VENTAS\CURSOS DE MI MODULO\DATOS PARA PROGRAMCION\MODELADO.jpg")
DEFAULT_EXCEL = Path(r"C:\Users\HP\Desktop\PROPUESTAS DE DIPLOMADO\DIPLOMADO EXCEL Y POWER BI APLICADO A LA GESTION DE VENTAS\CURSOS DE MI MODULO\DATOS PARA PROGRAMCION\Datos_Base.xlsx")
DEFAULT_SHEET = "Datos_base"
DEFAULT_MEASURES = Path(r"C:\Users\HP\Desktop\PROPUESTAS DE DIPLOMADO\DIPLOMADO EXCEL Y POWER BI APLICADO A LA GESTION DE VENTAS\CURSOS DE MI MODULO\DATOS PARA PROGRAMCION\ANALISIS BIVARIADO.txt")
DEFAULT_OUTPUT = Path(r"C:\Users\HP\Desktop\PROPUESTAS DE DIPLOMADO\DIPLOMADO EXCEL Y POWER BI APLICADO A LA GESTION DE VENTAS\CURSOS DE MI MODULO\DATOS PARA PROGRAMCION\output")


# Mapeo de nombre de medida a función de gráfica
PLOTTERS = {
    'ingresopromedioclienteproductoregion': 'plot_ingreso_promedio'
}

# ----------------------------------------
# Cliente OpenAI
# ----------------------------------------
def get_openai_client() -> OpenAI:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        logging.error("OPENAI_API_KEY no definida en el entorno")
        raise ValueError("Define OPENAI_API_KEY en las variables de entorno")
    return OpenAI(api_key=api_key)

# ----------------------------------------
# Utilidades de carga de datos
# ----------------------------------------
def load_image_data_uri(path: Path) -> str:
    raw = path.read_bytes()
    b64 = base64.b64encode(raw).decode()
    return f"data:image/png;base64,{b64}"


def load_excel(path: Path, sheet: str = None) -> pd.DataFrame:
    df = pd.read_excel(path, sheet_name=sheet)
    if isinstance(df.columns, pd.MultiIndex):
        df.columns = [" ".join(map(str, c)).strip() for c in df.columns]
    return df


def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    df = df.rename(columns=lambda c: c.strip())
    mapping = {}
    for col in df.columns:
        key = col.lower()
        if 'cliente' in key:
            mapping[col] = 'Num cliente'
        elif 'producto' in key:
            mapping[col] = 'Nombre del producto'
        elif any(term in key for term in ('departam','región','geografia','region')):
            mapping[col] = 'Departamento'
        elif 'ingres' in key:
            mapping[col] = 'Ingresos'
    return df.rename(columns=mapping)

# ----------------------------------------
# Snippet JSON para prompt con columnas únicas
# ----------------------------------------
def snippet_json(df: pd.DataFrame, n: int = 10) -> str:
    df2 = df.head(n).copy()
    # Asegurar nombres de columna únicos
    cols, counts = [], {}
    for col in df2.columns:
        cnt = counts.get(col, 0)
        new = col if cnt == 0 else f"{col}_{cnt}"
        cols.append(new)
        counts[col] = cnt + 1
    df2.columns = cols
    return json.dumps(df2.to_dict(orient='records'), indent=2, default=str)

# ----------------------------------------
# Carga de medidas DAX
# ----------------------------------------
def load_dax_measures(path: Path) -> list[str]:
    text = ''
    ext = path.suffix.lower()
    if ext == '.pdf':
        with pdfplumber.open(path) as pdf:
            text = '\n\n'.join(page.extract_text() for page in pdf.pages)
    else:
        text = path.read_text(encoding='utf-8')
    segments = [seg.strip() for seg in re.split(r'\n\s*\n', text) if seg.strip()]
    logging.info(f"Cargadas {len(segments)} medidas DAX desde {path}")
    return segments

# ----------------------------------------
# Llamada a ChatCompletion
# ----------------------------------------
def ask_openai(client: OpenAI, prompt: str) -> str:
    resp = client.chat.completions.create(
        model='gpt-4o-mini',
        messages=[{'role':'user', 'content': prompt}]
    )
    return resp.choices[0].message.content

# ----------------------------------------
# Funciones de cálculo y gráficos
# ----------------------------------------
def compute_ingreso_promedio(df: pd.DataFrame) -> pd.DataFrame:
    grouped = (
        df.groupby(['Num cliente', 'Nombre del producto', 'Departamento'])['Ingresos']
        .sum().reset_index(name='IngresoTotal')
    )
    avg = (
        grouped.groupby(['Departamento', 'Nombre del producto'])['IngresoTotal']
        .mean().reset_index(name='IngresoPromedio')
    )
    return avg


def plot_ingreso_promedio(df: pd.DataFrame, out_dir: Path):
    avg = compute_ingreso_promedio(df)
    csv_path = out_dir / 'ingreso_promedio.csv'
    avg.to_csv(csv_path, index=False)
    _plot_grouped_bar(
        avg, 'Departamento', 'Nombre del producto', 'IngresoPromedio',
        'Ingreso Promedio por Cliente', out_dir / 'ingreso_promedio.png'
    )


def _plot_grouped_bar(df: pd.DataFrame, index: str, column: str, value: str,
                      title: str, out_path: Path):
    pivot = df.pivot(index=index, columns=column, values=value)
    ax = pivot.plot(kind='bar')
    ax.set(title=title, xlabel=index, ylabel=value)
    plt.tight_layout()
    ax.figure.savefig(out_path)
    plt.close()
    logging.info(f"Guardado gráfico en {out_path}")

# ----------------------------------------
# Utilidades
# ----------------------------------------
def sanitize_measure_name(measure: str) -> str:
    name = measure.split('=')[0].strip()
    return re.sub(r'\W+', '', name).lower()

# ----------------------------------------
# Función principal
# ----------------------------------------
def main():
    parser = argparse.ArgumentParser(description='Analiza medidas DAX con OpenAI y genera gráficos')
    parser.add_argument('--image',    type=Path, default=DEFAULT_IMAGE)
    parser.add_argument('--excel',    type=Path, default=DEFAULT_EXCEL)
    parser.add_argument('--sheet',    default=DEFAULT_SHEET)
    parser.add_argument('--measures', type=Path, default=DEFAULT_MEASURES)
    parser.add_argument('--output',   type=Path, default=DEFAULT_OUTPUT)
    args, _ = parser.parse_known_args()

    args.output.mkdir(parents=True, exist_ok=True)
    out_docx = args.output / 'EXPLICACION.docx'

    client = get_openai_client()
    df = normalize_columns(load_excel(args.excel, args.sheet))
    measures = load_dax_measures(args.measures)

    doc = Document()
    doc.add_heading('Análisis de Medidas DAX', level=1)

    for m in measures:
        prompt = (
            f"Tengo este modelado de datos (imagen):\n"
            f"![Model]({load_image_data_uri(args.image)})\n\n"
            "Datos de muestra (primeras filas):\n```json\n"
            + snippet_json(df) + "\n```\n\n"
            "Analiza esta medida DAX y responde cada punto en orden:\n"
            f"```DAX\n{m}\n```\n"
            "1. ¿Para qué sirve?\n"
            "2. ¿Qué objeto visual de POWERBI con cenvenientes usar con esta medida?\n"
            "3. Ejemplo de uso de esta medida en un objeto visual adaptado a los datos que disponemos\n"
            "4. ¿Cómo interpretarlo?\n"
        )
        response = ask_openai(client, prompt)

        doc.add_heading(f'Medida: {m}', level=2)
        for line in response.split('\n'):
            doc.add_paragraph(line)

        key = sanitize_measure_name(m)
        plot_fn_name = PLOTTERS.get(key)
        if plot_fn_name:
            globals()[plot_fn_name](df, args.output)
            img_path = args.output / f'{key}.png'
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6))

    doc.save(out_docx)
    logging.info(f"Documento Word guardado en {out_docx}")

if __name__ == '__main__':
    main()


# In[ ]:




